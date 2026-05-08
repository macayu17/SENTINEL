from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_center(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_left(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def build_synopsis() -> None:
    doc = Document()

    add_center(doc, "VISVESVARAYA TECHNOLOGICAL UNIVERSITY")
    add_center(doc, '"JnanaSangama", Belagavi-590018.')
    add_left(doc, "")
    add_center(doc, "A Major Project (BCS685) Synopsis on")
    add_center(doc, '"SENTINEL: Smart Early-warning Network for Trading, Institutional Orders, and Liquidity Events"')
    add_left(doc, "Submitted in the partial fulfillment of the requirements for the award of the degree of")
    add_center(doc, "Bachelor of Engineering in Computer Science and Engineering")
    add_left(doc, "Submittedby")
    add_left(doc, "STUDENT NAME (USN)")
    add_left(doc, "STUDENT NAME (USN)")
    add_left(doc, "STUDENT NAME (USN)")
    add_left(doc, "STUDENT NAME (USN)")
    add_left(doc, "Project Team #")
    add_left(doc, "")
    add_left(doc, "Under the guidance of")
    add_left(doc, "Name")
    add_left(doc, "Designation,")
    add_left(doc, "Department of CS&E")
    add_left(doc, "RV INSTITUTE OF TECHNOLOGY AND MANAGEMENT")
    add_left(doc, "(Affiliated to Visvesvaraya Technological University, Belagavi & Approved by AICTE, New Delhi)")
    add_left(doc, "JP Nagar 8th Phase, Kothanur, Bengaluru-560076")
    add_center(doc, "2025-2026")

    add_left(doc, "")
    add_left(doc, "Abstract")
    add_left(
        doc,
        "SENTINEL is a real-time market microstructure simulator that models a multi-agent order book and "
        "generates early-warning signals for liquidity stress and large institutional orders. The system "
        "combines an event-driven simulation engine, predictive analytics, and a web dashboard to provide "
        "interpretable risk cues in sub-second time. It is designed for local execution and educational or "
        "research use, with optional RL-based policy control for advanced experimentation.",
    )

    add_left(doc, "")
    add_left(doc, "Background")
    add_left(
        doc,
        "Modern markets move rapidly, and conventional dashboards typically surface risk only after a "
        "liquidity event has already impacted price. A simulated, controllable order book allows researchers "
        "and traders to study microstructure dynamics and stress scenarios without live-market risk. SENTINEL "
        "fills this gap by combining agent-based simulation with live warnings and explainable metrics.",
    )

    add_left(doc, "")
    add_left(doc, "Objective")
    add_left(
        doc,
        "To build an early-warning system for liquidity shocks and large-order activity by simulating a "
        "multi-agent limit order book and streaming real-time risk indicators to a responsive dashboard.",
    )

    add_left(doc, "")
    add_left(doc, "Literature Survey (Tabular Format)")
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Author/Year"
    hdr[1].text = "Focus"
    hdr[2].text = "Key Findings"
    hdr[3].text = "Relevance to SENTINEL"

    rows = [
        ("Cont et al., 2011", "Order book modeling", "Shows importance of depth and spread in price impact", "Guides order book features"),
        ("Hasbrouck, 2007", "Market microstructure", "Explains liquidity measures and trade impact", "Motivates liquidity health score"),
        ("Gould et al., 2013", "Limit order books", "Reviews empirical LOB dynamics", "Informs simulation design"),
        ("Cartea et al., 2015", "Algorithmic trading", "Highlights agent behavior and latency effects", "Supports agent strategy variety"),
    ]

    for author, focus, findings, relevance in rows:
        row = table.add_row().cells
        row[0].text = author
        row[1].text = focus
        row[2].text = findings
        row[3].text = relevance

    add_left(doc, "")
    add_left(doc, "Methodology")
    add_left(
        doc,
        "The system runs an event-driven simulation loop that schedules agent wakeups and order arrivals "
        "based on latency. Each agent observes market state and submits orders to a price-time priority "
        "order book. Trades update market state and agent PnL, while prediction modules compute liquidity "
        "health and large-order signals. A FastAPI service exposes control endpoints and streams updates via "
        "WebSocket to a Next.js dashboard for visualization and alerting.",
    )

    add_left(doc, "")
    add_left(doc, "Hardware and Software Requirements")
    for item in [
        "CPU: 4-core laptop or better",
        "RAM: 8 GB or more",
        "OS: Windows, Linux, or macOS",
        "Python 3.10+ and Node.js 18+",
        "Docker (optional for full-stack run)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_left(doc, "")
    add_left(doc, "Expected outcome of the project")
    add_left(
        doc,
        "A working simulation environment with live dashboard that provides early warnings of liquidity "
        "stress and large-order activity, enabling safer experimentation with market microstructure and "
        "strategy behavior.",
    )

    add_left(doc, "")
    add_left(doc, "References")
    for ref in [
        "Cont, R. et al. (2011). The dynamics of order book markets.",
        "Hasbrouck, J. (2007). Empirical Market Microstructure.",
        "Gould, M. et al. (2013). Limit order books.",
        "Cartea, A. et al. (2015). Algorithmic and High-Frequency Trading.",
    ]:
        doc.add_paragraph(ref, style="List Number")

    output_path = r"d:\\Sentinel\\SENTINEL-main\\Synopsis_Sentinel.docx"
    doc.save(output_path)


if __name__ == "__main__":
    build_synopsis()
